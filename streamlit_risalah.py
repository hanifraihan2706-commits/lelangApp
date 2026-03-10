#!/usr/bin/env python3
"""
LELANG APPS — Web Edition (Streamlit)
Mengonversi data lelang dari Excel ke Risalah Lelang Word (.docx)

Author   : HANIF RAIHAN
Version  : 4.0 (Streamlit Web App)
Standard : OpenXML / python-docx  |  Kertas F4  |  Arial 12pt


# ══════════════════════════════════════════════════════════════════
#  IMPORTS
# ══════════════════════════════════════════════════════════════════
import io
import re
import os
import zipfile
from datetime import datetime

import streamlit as st
import openpyxl
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Cm, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PDF ringkasan statistik
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ══════════════════════════════════════════════════════════════════
#  CONSTANTS  (identik dengan versi desktop)
# ══════════════════════════════════════════════════════════════════
FONT_NAME       = "Arial"
FONT_SIZE_PT    = 12

PAGE_W_CM       = 21.59
PAGE_H_CM       = 33.02
MARGIN_L_CM     = 4.0
MARGIN_R_CM     = 1.5
MARGIN_T_CM     = 2.0
MARGIN_B_CM     = 2.0

CONTENT_W_CM    = PAGE_W_CM - MARGIN_L_CM - MARGIN_R_CM
TAB_TWIPS       = 600
TAB_CM          = TAB_TWIPS / 567
CONTENT_W_TWIPS = int(CONTENT_W_CM * 567)

POS_LABEL_L     = 600
POS_COLON       = 3600
POS_PEMBELI_L   = 600
POS_MID         = CONTENT_W_TWIPS // 2
TAB_DESC_START  = TAB_TWIPS
POS_RIGHT_CM    = 16.06
POS_RIGHT       = int(POS_RIGHT_CM * 567)

TAB_FULL        = POS_RIGHT
TAP_MID_POS     = POS_MID
TAP_END_POS     = POS_RIGHT
TAB_RIGHT_EDGE  = POS_RIGHT
TAB_LABEL_L     = POS_LABEL_L
TAB_COLON_L     = POS_COLON

POS_NILAI_L     = POS_LABEL_L
POS_NILAI_C     = POS_COLON
POS_HARGA_L     = POS_LABEL_L
POS_HARGA_C     = POS_COLON
POS_PEMB_L      = POS_LABEL_L
POS_PEMB_C      = POS_COLON

DASH_LINE       = "-" * 70

BULAN_ID = {
    "01": "Januari",  "02": "Februari", "03": "Maret",    "04": "April",
    "05": "Mei",      "06": "Juni",     "07": "Juli",     "08": "Agustus",
    "09": "September","10": "Oktober",  "11": "November", "12": "Desember"
}

ORDINAL_ID = [
    (1,"Pertama"),(2,"Kedua"),(3,"Ketiga"),(4,"Keempat"),(5,"Kelima"),
    (6,"Keenam"),(7,"Ketujuh"),(8,"Kedelapan"),(9,"Kesembilan"),(10,"Kesepuluh"),
    (11,"Kesebelas"),(12,"Keduabelas"),(13,"Ketigabelas"),(14,"Keempatbelas"),(15,"Kelimabelas"),
]

META_DEFAULTS = {
    "nomor_risalah"  : "027/10/19/2026",
    "tanggal_lelang" : "Rabu, delapan belas bulan Februari tahun dua ribu dua puluh enam (18-02-2026)",
    "pukul_lelang"   : "sebelas nol nol (11.00)",
    "nama_pejabat"   : "Rino Arief Rachman, S.H., M.H., M.Kn",
    "dasar_sk"       : "Surat Keputusan Menteri Keuangan Nomor 573/KM.6/2017 tanggal 20 Juni 2017",
    "wilayah_jabatan": "Kota Surabaya, Kabupaten Sidoarjo, dan sekitarnya",
    "alamat_pejabat" : "Jalan Perumahan Oasis Village Blok B-1, Semampir, Sedati, Kabupaten Sidoarjo, Provinsi Jawa Timur",
    "tempat_lelang"  : "Jalan Panglima Sudirman Nomor 17 B Menyanggong, Medaeng Taman, Kabupaten Sidoarjo, Jawa Timur",
    "nama_kuasa"     : "Adi Kurniawan Saputra",
    "jabatan_kuasa"  : "Branch Manager dan Kuasa Direksi",
    "nama_pt"        : "PT Anugerah Lelang Indonesia",
    "nomor_spl"      : "0075/SPL/ALI-SDA/II/2026",
    "tgl_spl"        : "11 Februari 2026",
}

COLUMN_ALIASES = {
    "lot"            : ["LOT", "ITEM LOT", "LOT ID", "NO LOT"],
    "nopol"          : ["NOPOL", "POLICE NUMBER", "NOMOR POLISI", "NO. POL", "PLAT", "NO POLISI"],
    "merk"           : ["MERK", "MEREK", "BRAND"],
    "model"          : ["MODEL", "TYPE", "TIPE", "MERK TYPE", "SERIES"],
    "no_mesin"       : ["MESIN", "NO MESIN", "MACHINE", "NOMOR MESIN", "MACHINE NUMBER"],
    "no_rangka"      : ["RANGKA", "NO RANGKA", "CHASSIS", "CHASSIS NUMBER", "NOMOR RANGKA"],
    "tahun"          : ["TAHUN", "YEAR"],
    "warna"          : ["WARNA", "COLOR"],
    "status"         : ["STATUS", "KETERANGAN"],
    "nama_customer"  : ["PEMBELI", "NAMA CUSTOMER", "CUSTOMER NAME", "PEMENANG"],
    "alamat_customer": ["ALAMAT", "CUSTOMER ADDRESS", "ALAMAT CUSTOMER", "ALAMAT PEMBELI"],
    "open_price"     : ["HARGA AWAL", "HARGA AWAL ALL", "BASIC PRICE ALL", "OPEN PRICE", "BASIC PRICE", "LIMIT", "NILAI LIMIT"],
    "harga_terbentuk": ["HARGA LELANG", "HARGA TERBENTUK", "AUCTION PRICE", "RESULT PRICE", "LAKU"],
}

# ══════════════════════════════════════════════════════════════════
#  UTILITY FUNCTIONS  (identik desktop)
# ══════════════════════════════════════════════════════════════════

def safe_str(val, default="-"):
    if val is None: return default
    if isinstance(val, float) and val == int(val): val = int(val)
    s = str(val).strip()
    return s if s and s.lower() not in ("none", "nan", "") else default

def format_rupiah(amount) -> str:
    try:
        if amount is None: return "Rp. 0,00"
        if isinstance(amount, str):
            amount = float(re.sub(r"[^\d,]", "", amount).replace(",", "."))
        amount = float(amount)
        if amount <= 0: return "Rp. 0,00"
        int_part = int(amount)
        dec_part = int(round((amount - int_part) * 100))
        return f"Rp. {'{:,}'.format(int_part).replace(',', '.')},{dec_part:02d}"
    except Exception:
        return "Rp. 0,00"

def _fmt_currency(amount: float) -> str:
    
      ≥ 1 Miliar  → "1,08 M"   (M = Miliar)
      ≥ 1 Juta    → "1.083 Jt" (Jt = Juta)
      < 1 Juta    → "950 rb"   (rb = ribu)
    """
    if amount is None or amount <= 0:
        return "Rp 0"
    if amount >= 1_000_000_000:
        return f"Rp {amount/1_000_000_000:.2f} M"
    if amount >= 1_000_000:
        return f"Rp {amount/1_000_000:,.0f} Jt"
    if amount >= 1_000:
        return f"Rp {amount/1_000:,.0f} rb"
    return f"Rp {amount:,.0f}"

def angka_ke_kata(n: int) -> str:
    satuan = ["","satu","dua","tiga","empat","lima","enam","tujuh","delapan","sembilan","sepuluh","sebelas"]
    if n == 0: return "nol"
    if n < 12: return satuan[n]
    if n < 20: return satuan[n-10] + " belas"
    if n < 100: return satuan[n//10] + " puluh" + (" " + satuan[n%10] if n%10 else "")
    if n < 1000:
        p = "se" if n//100 == 1 else satuan[n//100] + " "
        return p + "ratus" + (" " + angka_ke_kata(n%100) if n%100 else "")
    if n < 1_000_000:
        p = "se" if n//1000 == 1 else angka_ke_kata(n//1000) + " "
        return p + "ribu" + (" " + angka_ke_kata(n%1000) if n%1000 else "")
    return str(n)

def _extract_short_date(tanggal_full: str) -> str:
    m = re.search(r'\((\d{1,2})-(\d{2})-(\d{4})\)', tanggal_full)
    if m:
        d, mon, y = m.group(1), m.group(2), m.group(3)
        return f"{int(d)} {BULAN_ID.get(mon, mon)} {y}"
    m2 = re.search(r'(\d{2})-(\d{2})-(\d{4})', tanggal_full)
    if m2:
        d, mon, y = m2.group(1), m2.group(2), m2.group(3)
        return f"{int(d)} {BULAN_ID.get(mon, mon)} {y}"
    return tanggal_full

# ══════════════════════════════════════════════════════════════════
#  WORD XML HELPERS  (identik desktop — kunci kerapian dokumen)
# ══════════════════════════════════════════════════════════════════

def _set_run_font(run, size_pt=FONT_SIZE_PT, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)

def _add_tabstops(para, stops: list):
    _LMAP = {"dash":"hyphen","dot":"dot","line":"underscore"}
    pPr = para._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = OxmlElement("w:tabs")
        pPr.append(tabs_el)
    for stop in stops:
        pos    = stop[0]
        leader = stop[1] if len(stop) > 1 else None
        align  = stop[2] if len(stop) > 2 else "left"
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), align)
        tab.set(qn("w:pos"), str(int(pos)))
        if leader is not None:
            tab.set(qn("w:leader"), _LMAP.get(leader, "hyphen"))
        tabs_el.append(tab)

def _new_para(doc, left_cm=0.0, first_cm=0.0, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before_pt=0, space_after_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.left_indent       = Cm(left_cm)
    pf.first_line_indent = Cm(first_cm)
    pf.space_before      = Pt(space_before_pt)
    pf.space_after       = Pt(space_after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p

def _add_run(para, text, **kwargs):
    run = para.add_run(text)
    _set_run_font(run, **kwargs)
    return run

def _make_xml_run(text=None, is_instr=False, fld_char_type=None,
                  size_half_pt=20, bold=False, italic=False):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for attr in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rf.set(qn(attr), FONT_NAME)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size_half_pt))
    rPr.append(rf); rPr.append(sz)
    if bold:   rPr.append(OxmlElement("w:b"))
    if italic: rPr.append(OxmlElement("w:i"))
    r.append(rPr)
    if fld_char_type is not None:
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), fld_char_type); r.append(fc)
    elif is_instr:
        it = OxmlElement("w:instrText")
        it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        it.text = str(text); r.append(it)
    else:
        t = OxmlElement("w:t")
        s = str(text) if text is not None else ""
        if s.startswith(" ") or s.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = s; r.append(t)
    return r

def _make_page_field_runs(**kwargs):
    return [
        _make_xml_run(fld_char_type="begin", **kwargs),
        _make_xml_run(" PAGE ", is_instr=True, **kwargs),
        _make_xml_run(fld_char_type="separate", **kwargs),
        _make_xml_run("1", **kwargs),
        _make_xml_run(fld_char_type="end", **kwargs),
    ]

def _build_nested_if_ordinal(parent_el, ordinal_list, **run_kwargs):
    if not ordinal_list: return
    page_num, ordinal = ordinal_list[0]
    remaining = ordinal_list[1:]
    parent_el.append(_make_xml_run(fld_char_type="begin", **run_kwargs))
    parent_el.append(_make_xml_run(" IF ", is_instr=True, **run_kwargs))
    for r in _make_page_field_runs(**run_kwargs): parent_el.append(r)
    if remaining:
        parent_el.append(_make_xml_run(f' = {page_num} "{ordinal}" ', is_instr=True, **run_kwargs))
        _build_nested_if_ordinal(parent_el, remaining, **run_kwargs)
    else:
        parent_el.append(_make_xml_run(f' = {page_num} "{ordinal}" "Berikutnya" ', is_instr=True, **run_kwargs))
    parent_el.append(_make_xml_run(fld_char_type="separate", **run_kwargs))
    parent_el.append(_make_xml_run(ordinal, **run_kwargs))
    parent_el.append(_make_xml_run(fld_char_type="end", **run_kwargs))

def _add_footer(doc, meta: dict):
    """
    PATCH v4.1 — Pindah ke HEADER (atas kertas) sesuai format resmi:
      BEFORE: section.footer  → teks muncul di bawah halaman
      AFTER:  section.header  → teks muncul di atas setiap halaman

    Susunan header per halaman:
      Lembar [Ordinal] dari Risalah Lelang Nomor X tanggal Y   |   Pejabat Lelang Kelas II,
      ─────────────────────────────────────────────────────────      [kosong]
                                                                     ttd
                                                                     [Nama Pejabat]
    """
    nomor     = meta.get("nomor_risalah", "027/10/19/2026")
    tgl_raw   = meta.get("tanggal_lelang", "(18-02-2026)")
    tgl_short = _extract_short_date(tgl_raw)
    pejabat   = meta.get("nama_pejabat", "Rino Arief Rachman, S.H., M.H., M.Kn")

    section = doc.sections[0]

    # ── BEFORE: section.footer  |  AFTER: section.header ─────────
    header = section.header
    hd     = header._element          # <w:hdr> element
    for p_elem in list(hd.findall(qn("w:p"))): hd.remove(p_elem)

    def _new_hd_para(align="left", sp_before=0, sp_after=0):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(int(sp_before * 20)))
        sp.set(qn("w:after"),  str(int(sp_after  * 20)))
        sp.set(qn("w:line"), "240"); sp.set(qn("w:lineRule"), "auto")
        pPr.append(sp)
        if align != "left":
            jc = OxmlElement("w:jc"); jc.set(qn("w:val"), align); pPr.append(jc)
        p.append(pPr); hd.append(p); return p

    SIZE = 20   # 10pt (half-points)

    # ── Baris 1: "Lembar [ORDINAL] dari Risalah Lelang ..." ───────
    p_lembar = _new_hd_para(sp_before=0, sp_after=0)
    p_lembar.append(_make_xml_run("Lembar ", size_half_pt=SIZE))
    _build_nested_if_ordinal(p_lembar, ORDINAL_ID, size_half_pt=SIZE)
    p_lembar.append(_make_xml_run(
        f" dari Risalah Lelang Nomor {nomor} tanggal {tgl_short}",
        size_half_pt=SIZE))

    # ── Baris 2: Garis pemisah ─────────────────────────────────────
    p_sep = _new_hd_para(sp_before=0, sp_after=2)
    p_sep.append(_make_xml_run(DASH_LINE, size_half_pt=SIZE))

    # ── Blok TTD rata kanan ────────────────────────────────────────
    for text, bold, italic in [
        ("Pejabat Lelang Kelas II,", False, False),
        ("",                          False, False),
        ("ttd",                       False, True),
        ("",                          False, False),
        (pejabat,                     True,  False),
    ]:
        p = _new_hd_para(align="right", sp_before=0, sp_after=0)
        p.append(_make_xml_run(text, size_half_pt=SIZE, bold=bold, italic=italic))

# ══════════════════════════════════════════════════════════════════
#  LOT ENTRY BUILDERS  (identik desktop)
# ══════════════════════════════════════════════════════════════════

def _build_description(item: dict) -> str:
    jenis = "roda dua" if item.get("jenis") == "motor" else "roda empat"
    return (f"Satu (1) unit kendaraan {jenis}, Merk {item['merk']}, "
            f"Model {item['model']}, Tahun {item['tahun']}, Warna {item['warna']}, "
            f"Nomor Polisi {item['nopol']}, Nomor Rangka")

def _line1(doc, txt, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = _new_para(doc, left_cm=TAB_CM, first_cm=-TAB_CM, align=align)
    _add_tabstops(p, [(TAB_DESC_START, None, "left"), (POS_RIGHT, "dash", "right")])
    _add_run(p, f"{txt}\t"); return p

def _label_line_triple(doc, label, value, align=WD_ALIGN_PARAGRAPH.LEFT):
    target = POS_COLON + 240
    p = _new_para(doc, left_cm=target/567, first_cm=(POS_LABEL_L-target)/567, align=align)
    _add_tabstops(p, [(POS_COLON, "dash", "left"), (POS_RIGHT, "dash", "right")])
    _add_run(p, f"{label}\t: {value}\t"); return p

def _label_pembeli(doc, label, value, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = _new_para(doc, left_cm=TAB_TWIPS/567, first_cm=(POS_PEMBELI_L-TAB_TWIPS)/567, align=align)
    _add_tabstops(p, [(POS_COLON, "dash", "left"), (POS_RIGHT, "dash", "right")])
    _add_run(p, f"{label}\t: {value}\t"); return p

def add_lot_entry(doc, item: dict, section: str, seq_num: int = None):
    lot_num   = item.get("lot", 0)
    no_rangka = item.get("no_rangka", "-")
    no_mesin  = item.get("no_mesin",  "-")
    desc      = _build_description(item)
    list_num  = seq_num if seq_num is not None else lot_num
    lot_pfx   = f"Lot {lot_num}, " if section in ("SOLD", "TAP") else ""

    if section == "LIMIT":
        val = format_rupiah(item.get("limit", 0))
        _line1(doc, f"{list_num}.\t{desc} {no_rangka}, Nomor Mesin {no_mesin},")
        _label_line_triple(doc, "Nilai Limit", f"{val}.")

    elif section == "SOLD":
        val     = format_rupiah(item.get("harga_terbentuk", 0))
        pembeli = f"{item.get('nama_customer','-')}, beralamat di {item.get('alamat_customer','-')}"
        _line1(doc, f"{list_num}.\t{lot_pfx}{desc} {no_rangka}, Nomor Mesin {no_mesin},")
        _label_line_triple(doc, "Harga Lelang", f"{val}.")
        _label_pembeli(doc, "Pembeli", f"{pembeli}.")

    elif section == "TAP":
        _line1(doc, f"{list_num}.\t{lot_pfx}{desc} {no_rangka}, Nomor Mesin {no_mesin}.",
               align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        p = _new_para(doc, left_cm=TAB_CM, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        mid = int(POS_RIGHT / 2)
        _add_tabstops(p, [(mid, "dash", "center"), (POS_RIGHT, "dash", "right")])
        _add_run(p, "\tTAP\t")

# ══════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS  (identik desktop)
# ══════════════════════════════════════════════════════════════════

def _set_page_format(doc):
    sec = doc.sections[0]
    sec.page_width    = Cm(21.59)
    sec.page_height   = Cm(33.02)
    sec.left_margin   = Cm(4.0)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

def _clear_default_para(doc):
    for para in doc.paragraphs[:]:
        para._element.getparent().remove(para._element)

def _add_header_section(doc, meta, items):
    def center_bold(text, size=12, extra_space=0):
        p = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=extra_space)
        _add_run(p, text, size_pt=size, bold=True)
    def justify_text(text, sp_after=3):
        p = _new_para(doc, space_after_pt=sp_after); _add_run(p, text)
    def separator():
        p = _new_para(doc, space_after_pt=0); _add_run(p, DASH_LINE)

    center_bold("RISALAH LELANG", size=14)
    center_bold(meta.get("nomor_risalah", "-"), size=12, extra_space=4)
    separator()
    justify_text(
        f"----Pada hari ini, {meta.get('tanggal_lelang','-')}, "
        f"dimulai pukul {meta.get('pukul_lelang','-')} Waktu Server, "
        f"di hadapan Saya {meta.get('nama_pejabat','-')}, Pejabat Lelang Kelas II "
        f"yang diangkat berdasarkan {meta.get('dasar_sk','-')}, "
        f"dengan wilayah jabatan meliputi {meta.get('wilayah_jabatan','-')}, "
        f"berkedudukan di {meta.get('alamat_pejabat','-')}, "
        f"dilaksanakan Lelang Sukarela atas Barang Bergerak Kendaraan Bermotor "
        f"yang akan diuraikan lebih lanjut di bawah ini.", sp_after=0)
    separator()
    justify_text(
        f"----Pelaksanaan lelang ini dilakukan atas permintaan Saudara "
        f"{meta.get('nama_kuasa','-')}, selaku {meta.get('jabatan_kuasa','-')} "
        f"{meta.get('nama_pt','-')}, sesuai Surat Permohonan Lelang Nomor: "
        f"{meta.get('nomor_spl','-')} tanggal {meta.get('tgl_spl','-')}.", sp_after=0)
    separator()
    jml = len(items); jml_w = angka_ke_kata(jml)
    jenis_barang = ("kendaraan bermotor roda dua"
                    if any(i["jenis"] == "motor" for i in items)
                    else "kendaraan bermotor roda empat")
    justify_text(
        f"----Barang bergerak yang dilelang berupa {jml} ({jml_w}) lot {jenis_barang} "
        f"yang terdiri dari berbagai macam jenis merk, type, sesuai dengan yang tercantum "
        f"dalam daftar barang yang telah dibagikan kepada peserta lelang, yaitu:", sp_after=4)
    separator()

def _add_closing_section(doc, meta, items):
    def justify_text(text, sp_after=0, sp_before=0):
        p = _new_para(doc, space_after_pt=sp_after, space_before_pt=sp_before)
        _add_run(p, text)

    pejabat   = meta.get("nama_pejabat", "-")
    sold_items = [i for i in items if i["status"].upper() in ("SOLD","TERJUAL","LAKU")]
    tap_items  = [i for i in items if i["status"].upper() not in ("SOLD","TERJUAL","LAKU")]

    justify_text(
        f"----Dengan demikian berakhirlah pelaksanaan lelang pada hari ini. "
        f"Hasil lelang: {len(sold_items)} ({angka_ke_kata(len(sold_items))}) lot terjual/laku. "
        f"{len(tap_items)} ({angka_ke_kata(len(tap_items))}) lot tidak terjual/ditarik. "
        f"Risalah Lelang ini dibuat sebagai bukti pelaksanaan lelang.",
        sp_after=4, sp_before=10)
    p_sep = _new_para(doc, space_after_pt=0); _add_run(p_sep, DASH_LINE)
    p_ttd = _new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before_pt=36, space_after_pt=48)
    _add_run(p_ttd, "Pejabat Lelang Kelas II,")
    p_name = _new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after_pt=0)
    _add_run(p_name, pejabat, bold=True)

# ══════════════════════════════════════════════════════════════════
#  EXCEL DATA READER
#  PERUBAHAN: terima BytesIO (upload Streamlit) selain path string
# ══════════════════════════════════════════════════════════════════

def _find_header_and_map(ws, log_list):
    header_idx = -1
    col_map    = {}
    for idx, row in enumerate(ws.iter_rows(min_row=1, max_row=15, values_only=True), start=1):
        row_strs = [str(c).strip().upper() for c in row if c is not None]
        if any("LOT" in s for s in row_strs):
            header_idx = idx
            for col_i, cell in enumerate(row):
                if cell is None: continue
                cell_val = str(cell).strip().upper()
                for key, aliases in COLUMN_ALIASES.items():
                    if any(alias in cell_val for alias in aliases):
                        if key == "nopol":
                            if key not in col_map: col_map[key] = []
                            col_map[key].append(col_i)
                        elif key not in col_map:
                            col_map[key] = col_i
            break
    if header_idx != -1:
        log_list.append(f"🔍 Header ditemukan di baris ke-{header_idx}")
    return header_idx, col_map

def read_excel_data(source, sheet_name: str, log_list=None) -> list:
    """source: str (path) atau BytesIO (upload Streamlit)"""
    if log_list is None: log_list = []

    if isinstance(source, (str, bytes, os.PathLike)):
        wb = openpyxl.load_workbook(source, data_only=True)
    else:
        wb = openpyxl.load_workbook(io.BytesIO(source.read()) if hasattr(source, "read") else source, data_only=True)

    ws   = wb[sheet_name]
    jenis = "motor" if "MOTOR" in sheet_name.upper() else "mobil"

    header_idx, col_map = _find_header_and_map(ws, log_list)
    if header_idx == -1 or "lot" not in col_map:
        log_list.append("❌ Gagal mendeteksi header. Pastikan ada kolom 'LOT'.")
        return []

    items = []
    for row in ws.iter_rows(min_row=header_idx + 1, values_only=True):
        def _get(col_key):
            idx = col_map.get(col_key, -1)
            if isinstance(idx, list):
                parts = [str(row[i]).strip() for i in idx if i < len(row) and row[i] is not None]
                return " ".join(parts) if parts else "-"
            if idx < 0 or idx >= len(row): return None
            return row[idx]

        def _get_numeric(col_key):
            val = _get(col_key)
            if val is None or val == "": return 0
            if isinstance(val, (int, float)): return float(val)
            try:
                return float(str(val).replace("Rp","").replace(".","").replace(",",".").strip())
            except: return 0

        lot_val = _get("lot")
        if lot_val is None: continue
        try:   lot_num = int(float(lot_val))
        except: continue

        no_mesin_raw = _get("no_mesin")
        if isinstance(no_mesin_raw, (float, int)):
            no_mesin_raw = str(int(no_mesin_raw))

        item = {
            "lot"            : lot_num,
            "nopol"          : safe_str(_get("nopol")),
            "merk"           : safe_str(_get("merk")),
            "model"          : safe_str(_get("model")),
            "no_mesin"       : safe_str(no_mesin_raw),
            "no_rangka"      : safe_str(_get("no_rangka")),
            "tahun"          : safe_str(_get("tahun")),
            "warna"          : safe_str(_get("warna")),
            "status"         : safe_str(_get("status"), "Not Sold").strip().upper(),
            "nama_customer"  : safe_str(_get("nama_customer")),
            "alamat_customer": safe_str(_get("alamat_customer")),
            "limit"          : _get_numeric("open_price"),
            "harga_terbentuk": _get_numeric("harga_terbentuk"),
            "jenis"          : jenis,
        }
        items.append(item)

    log_list.append(f"✅ Total lot terbaca: {len(items)}")
    return items

def resolve_section(status: str) -> str:
    s = status.upper().replace(" ", "")
    if s in ("SOLD","TERJUAL","LAKU"):  return "SOLD"
    if s in ("NOTSOLD","TAKTERJUAL","BELUMLAKU",""): return "LIMIT"
    return "TAP"

# ══════════════════════════════════════════════════════════════════
#  DOCUMENT GENERATOR
#  PERUBAHAN: simpan ke BytesIO (bukan disk), terima BytesIO source
# ══════════════════════════════════════════════════════════════════

def generate_doc_bytes(source, sheet_name: str, meta: dict) -> tuple:
    """
    Generate dokumen Word dan kembalikan bytes-nya.
    Returns: (True, bytes) atau (False, error_str)
    """
    log_list = []
    try:
        # Reset cursor jika BytesIO
        if hasattr(source, "seek"): source.seek(0)

        wb_test = openpyxl.load_workbook(
            io.BytesIO(source.read()) if hasattr(source, "read") else source,
            read_only=True)
        if sheet_name not in wb_test.sheetnames:
            return False, f"Sheet '{sheet_name}' tidak ditemukan."
        wb_test.close()

        if hasattr(source, "seek"): source.seek(0)
        items = read_excel_data(source, sheet_name, log_list)
        if not items:
            return False, "Tidak ada data lot yang terbaca."

        STATUS_LAKU = ("SOLD","TERJUAL","LAKU")
        sold_items = [i for i in items if i["status"] in STATUS_LAKU]
        tap_items  = [i for i in items if i["status"] not in STATUS_LAKU]

        doc = Document()
        _set_page_format(doc)
        _clear_default_para(doc)
        doc.styles["Normal"].font.name = FONT_NAME
        doc.styles["Normal"].font.size = Pt(FONT_SIZE_PT)
        _clear_default_para(doc)

        _add_footer(doc, meta)
        _add_header_section(doc, meta, items)

        def _sec_hdr(title):
            p = _new_para(doc, space_before_pt=12, space_after_pt=6)
            p.add_run(title).bold = True

        # Seksi 1 — Semua Lot (Nilai Limit)
        for seq, itm in enumerate(items, 1):
            try: add_lot_entry(doc, itm, "LIMIT", seq_num=seq)
            except: pass

        # Seksi 2 — SOLD
        if sold_items:
            p = _new_para(doc, space_before_pt=12, space_after_pt=0); _add_run(p, "-"*70)
            _sec_hdr("Barang laku terjual:")
            p2 = _new_para(doc, space_after_pt=6); _add_run(p2, "-"*70)
            for seq, itm in enumerate(sold_items, 1):
                try: add_lot_entry(doc, itm, "SOLD", seq_num=seq)
                except: pass

        # Seksi 3 — TAP
        if tap_items:
            p = _new_para(doc, space_before_pt=12, space_after_pt=0); _add_run(p, "-"*70)
            _sec_hdr("Barang tidak laku terjual:")
            p2 = _new_para(doc, space_after_pt=6); _add_run(p2, "-"*70)
            for seq, itm in enumerate(tap_items, 1):
                try: add_lot_entry(doc, itm, "TAP", seq_num=seq)
                except: pass

        _add_closing_section(doc, meta, items)

        # ── Simpan ke BytesIO ─────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return True, buf.getvalue(), items

    except Exception as exc:
        import traceback
        return False, f"{exc}\n{traceback.format_exc()}", []

# ══════════════════════════════════════════════════════════════════
#  STATISTIK — CHART & PDF
# ══════════════════════════════════════════════════════════════════

def build_stat_charts(items: list, filename: str) -> tuple:
    """
    Buat 2 chart matplotlib:
      1. Pie chart: SOLD vs TAP vs LIMIT (not sold)
      2. Bar chart: Total Limit vs Total Harga Terbentuk
    Returns: (fig_pie, fig_bar, stat_dict)
    """
    STATUS_LAKU = ("SOLD","TERJUAL","LAKU")
    sold  = [i for i in items if i["status"] in STATUS_LAKU]
    tap   = [i for i in items if i["status"] not in STATUS_LAKU and resolve_section(i["status"]) == "TAP"]
    limit = [i for i in items if resolve_section(i["status"]) == "LIMIT"]

    n_sold  = len(sold)
    n_tap   = len(tap)
    n_limit = len(limit)

    total_limit    = sum(i.get("limit", 0) or 0 for i in items)
    total_harga    = sum(i.get("harga_terbentuk", 0) or 0 for i in sold)
    pencapaian_pct = (total_harga / total_limit * 100) if total_limit > 0 else 0

    stat = {
        "total": len(items),
        "sold" : n_sold,
        "tap"  : n_tap,
        "limit": n_limit,
        "total_limit": total_limit,
        "total_harga": total_harga,
        "pencapaian" : pencapaian_pct,
    }

    # ─── Pie Chart ────────────────────────────────────────────────
    palette = {"SOLD":"#22c55e", "TAP":"#ef4444", "LIMIT":"#f59e0b"}
    sizes, labels, pie_colors = [], [], []
    for key, n, lbl in [("SOLD", n_sold, f"SOLD ({n_sold})"),
                         ("TAP",  n_tap,  f"TAP ({n_tap})"),
                         ("LIMIT",n_limit,f"NOT SOLD ({n_limit})")]:
        if n > 0:
            sizes.append(n); labels.append(lbl); pie_colors.append(palette[key])

    fig_pie, ax_pie = plt.subplots(figsize=(5, 4), facecolor="#0f172a")
    ax_pie.set_facecolor("#0f172a")
    if sizes:
        wedges, texts, autotexts = ax_pie.pie(
            sizes, labels=labels, colors=pie_colors, autopct="%1.1f%%",
            startangle=90, pctdistance=0.75,
            wedgeprops=dict(width=0.55, edgecolor="#0f172a", linewidth=2))
        for t in texts:     t.set_color("#e2e8f0"); t.set_fontsize(9)
        for t in autotexts: t.set_color("white");   t.set_fontsize(8); t.set_fontweight("bold")
    ax_pie.set_title(f"Komposisi Hasil Lelang\n{filename}",
                     color="#e2e8f0", fontsize=10, pad=10)
    plt.tight_layout()

    # ─── Bar Chart ────────────────────────────────────────────────
    fig_bar, ax_bar = plt.subplots(figsize=(5, 4), facecolor="#0f172a")
    ax_bar.set_facecolor("#1e293b")
    bars = ax_bar.bar(
        ["Total Limit", "Harga Terbentuk"],
        [total_limit / 1e6, total_harga / 1e6],
        color=["#3b82f6", "#22c55e"], width=0.45,
        edgecolor="#0f172a", linewidth=1.5)
    for bar, val in zip(bars, [total_limit, total_harga]):
        ax_bar.text(bar.get_x() + bar.get_width()/2,
                    bar.get_height() + max(total_limit, total_harga or 1) / 1e6 * 0.02,
                    f"Rp {val/1e6:.1f}M",
                    ha="center", va="bottom", color="#e2e8f0", fontsize=9, fontweight="bold")
    ax_bar.set_ylabel("Nilai (Juta Rp)", color="#94a3b8", fontsize=9)
    ax_bar.tick_params(colors="#94a3b8")
    ax_bar.spines[:].set_color("#334155")
    ax_bar.set_title(f"Pencapaian Harga — {pencapaian_pct:.1f}%",
                     color="#e2e8f0", fontsize=10, pad=10)
    for spine in ax_bar.spines.values(): spine.set_color("#334155")
    ax_bar.set_facecolor("#1e293b")
    ax_bar.yaxis.label.set_color("#94a3b8")
    ax_bar.tick_params(axis="x", colors="#e2e8f0")
    ax_bar.tick_params(axis="y", colors="#94a3b8")
    plt.tight_layout()

    return fig_pie, fig_bar, stat


def build_stat_pdf(items: list, meta: dict, filename: str,
                   fig_pie: plt.Figure, fig_bar: plt.Figure) -> bytes:
    """Ekspor ringkasan statistik lelang ke PDF (1 halaman A4)."""
    buf     = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buf, pagesize=A4,
                                topMargin=1.5*cm, bottomMargin=1.5*cm,
                                leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []

    # ── Header ──────────────────────────────────────────────────
    style_title = ParagraphStyle("title", parent=styles["Heading1"],
                                  alignment=TA_CENTER, fontSize=14,
                                  textColor=colors.HexColor("#1e3a5f"), spaceAfter=4)
    style_sub   = ParagraphStyle("sub", parent=styles["Normal"],
                                  alignment=TA_CENTER, fontSize=10,
                                  textColor=colors.grey, spaceAfter=12)
    style_body  = ParagraphStyle("body", parent=styles["Normal"], fontSize=9, spaceAfter=3)

    story.append(Paragraph("LAPORAN PERFORMA LELANG", style_title))
    story.append(Paragraph(f"Nomor Risalah: {meta.get('nomor_risalah','-')}", style_sub))
    story.append(Paragraph(f"Tanggal       : {meta.get('tanggal_lelang','-')[:50]}", style_sub))
    story.append(Paragraph(f"Pejabat Lelang: {meta.get('nama_pejabat','-')}", style_sub))
    story.append(Spacer(1, 8))

    # ── Statistik Tabel ─────────────────────────────────────────
    STATUS_LAKU = ("SOLD","TERJUAL","LAKU")
    sold_items = [i for i in items if i["status"] in STATUS_LAKU]
    tap_items  = [i for i in items if i["status"] not in STATUS_LAKU]
    total_limit = sum(i.get("limit", 0) or 0 for i in items)
    total_harga = sum(i.get("harga_terbentuk", 0) or 0 for i in sold_items)
    pct = (total_harga / total_limit * 100) if total_limit else 0

    tbl_data = [
        ["Keterangan", "Nilai"],
        ["Total Lot",            str(len(items))],
        ["Lot Terjual (SOLD)",   str(len(sold_items))],
        ["Lot Tidak Terjual",    str(len(tap_items))],
        ["Total Nilai Limit",    format_rupiah(total_limit)],
        ["Total Harga Terbentuk",format_rupiah(total_harga)],
        ["Pencapaian Harga",     f"{pct:.1f}%"],
    ]
    tbl = Table(tbl_data, colWidths=[9*cm, 8*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR",   (0,0), (-1,0),  colors.white),
        ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#f0f4f8"), colors.white]),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.HexColor("#cccccc")),
        ("ALIGN",       (1,0), (1,-1),  "RIGHT"),
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",  (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0),(-1,-1), 4),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 12))

    # ── Charts ──────────────────────────────────────────────────
    def fig_to_img(fig):
        img_buf = io.BytesIO()
        fig.savefig(img_buf, format="png", dpi=120, bbox_inches="tight",
                    facecolor=fig.get_facecolor())
        img_buf.seek(0)
        return img_buf

    img_pie = fig_to_img(fig_pie)
    img_bar = fig_to_img(fig_bar)

    chart_table = Table(
        [[Image(img_pie, width=8.5*cm, height=7*cm),
          Image(img_bar, width=8.5*cm, height=7*cm)]],
        colWidths=[9*cm, 9*cm])
    chart_table.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"TOP")]))
    story.append(chart_table)
    story.append(Spacer(1, 12))

    # ── Footer ──────────────────────────────────────────────────
    story.append(Paragraph(
        f"Laporan dibuat otomatis oleh LELANG APPS v4.0  |  "
        f"{datetime.now().strftime('%d %B %Y  %H:%M')}",
        ParagraphStyle("footer", parent=styles["Normal"],
                       fontSize=7, textColor=colors.grey, alignment=TA_CENTER)))

    doc_pdf.build(story)
    buf.seek(0)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════
#  STREAMLIT UI
# ══════════════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="Lelang Apps — Risalah Lelang Generator",
        page_icon="⚖️",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # ─── Custom CSS ──────────────────────────────────────────────
    st.markdown("""
    <style>
    .stApp { background-color: #0f172a; color: #e2e8f0; }
    .block-container { padding-top: 1.5rem; }
    .stTabs [data-baseweb="tab-list"] { background: #1e293b; border-radius: 8px; padding: 4px; }
    .stTabs [data-baseweb="tab"] { color: #94a3b8; }
    .stTabs [aria-selected="true"] { background: #3b82f6 !important; color: white !important; border-radius: 6px; }
    /* Metric widget override — teks tidak terpotong */
    div[data-testid="stMetric"] {
        background: #1e293b; border: 1px solid #334155;
        border-radius: 12px; padding: 14px 10px;
    }
    div[data-testid="stMetricValue"] > div {
        color: #e2e8f0; font-size: 1.35rem !important;
        word-break: break-word; white-space: normal !important;
        line-height: 1.2;
    }
    div[data-testid="stMetricLabel"] > div {
        color: #94a3b8; font-size: 0.72rem;
        text-transform: uppercase; letter-spacing: .04em;
    }
    div[data-testid="stMetricDelta"] > div { font-size: 0.72rem; }
    .stDownloadButton button { width: 100%; }
    </style>
    """, unsafe_allow_html=True)

    # ─── Top Bar ─────────────────────────────────────────────────
    st.markdown("""
    <div style="background:#1e293b;padding:16px 24px;border-radius:12px;margin-bottom:20px;
                border-left:4px solid #3b82f6;display:flex;align-items:center;gap:16px">
      <span style="font-size:2rem">⚖️</span>
      <div>
        <div style="font-size:1.4rem;font-weight:800;color:#e2e8f0">LELANG APPS — PL2RINO</div>
        <div style="font-size:0.8rem;color:#94a3b8">
            Risalah Lelang Generator v4.0   |  #by.haniffraihan
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ─── Sidebar: Upload & Meta ───────────────────────────────────
    with st.sidebar:
        st.markdown("### 📁 Upload & Konfigurasi")

        uploaded_files = st.file_uploader(
            "Upload File Excel (.xlsx)",
            type=["xlsx", "xls"],
            accept_multiple_files=True,
            help="Ctrl+Klik untuk memilih banyak file sekaligus")

        sheet_name = "FIRMAN"
        if uploaded_files:
            try:
                first_bytes = uploaded_files[0].read()
                uploaded_files[0].seek(0)
                wb_tmp = openpyxl.load_workbook(io.BytesIO(first_bytes), read_only=True)
                sheets = wb_tmp.sheetnames
                wb_tmp.close()
                sheet_name = st.selectbox("🗂️ Pilih Sheet", sheets,
                                           index=0 if "FIRMAN" not in sheets
                                           else sheets.index("FIRMAN"))
            except Exception as e:
                st.warning(f"Gagal baca sheet: {e}")

        st.divider()
        st.markdown("### 📋 Data Dokumen")

        meta = {}
        field_labels = [
            ("nomor_risalah",   "Nomor Risalah"),
            ("tanggal_lelang",  "Tanggal Lelang"),
            ("pukul_lelang",    "Pukul Lelang"),
            ("nama_pejabat",    "Nama Pejabat Lelang"),
            ("dasar_sk",        "Dasar SK Pejabat"),
            ("wilayah_jabatan", "Wilayah Jabatan"),
            ("alamat_pejabat",  "Alamat Pejabat"),
            ("tempat_lelang",   "Tempat Lelang"),
            ("nama_kuasa",      "Nama Kuasa Penjual"),
            ("jabatan_kuasa",   "Jabatan Kuasa"),
            ("nama_pt",         "Nama PT / Penjual"),
            ("nomor_spl",       "Nomor SPL"),
            ("tgl_spl",         "Tanggal SPL"),
        ]
        for key, label in field_labels:
            meta[key] = st.text_input(label, value=META_DEFAULTS.get(key, ""), key=f"meta_{key}")

    # ─── TABS ────────────────────────────────────────────────────
    tab_gen, tab_stat = st.tabs(["⚡ Generate Risalah", "📊 Statistik Lelang"])

    # ╔══════════════════════════════════════════════════════════════
    # ║  TAB 1 — GENERATE RISALAH
    # ╚══════════════════════════════════════════════════════════════
    with tab_gen:
        if not uploaded_files:
            st.info("👆 Upload satu atau lebih file Excel dari sidebar untuk mulai.", icon="📂")
            st.stop()

        st.markdown(f"**{len(uploaded_files)} file siap diproses pada sheet `{sheet_name}`**")

        col_btn, col_space = st.columns([1, 3])
        with col_btn:
            do_generate = st.button("⚡ Generate Semua Risalah",
                                    type="primary", use_container_width=True)

        if do_generate:
            results  = []
            errors   = []
            all_items_map = {}
            progress = st.progress(0, text="Memulai proses batch...")
            log_area = st.expander("📋 Log Proses", expanded=True)
            log_lines = []

            for idx, upf in enumerate(uploaded_files, 1):
                nama = upf.name
                progress.progress(int((idx-1)/len(uploaded_files)*100),
                                   text=f"Memproses {idx}/{len(uploaded_files)}: {nama}")
                upf.seek(0)
                ok, result, items = generate_doc_bytes(upf, sheet_name, meta)

                if ok:
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    out_name = f"RISALAH_{os.path.splitext(nama)[0]}_{ts}.docx"
                    results.append((out_name, result))
                    all_items_map[nama] = items
                    log_lines.append(f"✅ [{idx}] {nama} → {out_name}")
                else:
                    errors.append(nama)
                    log_lines.append(f"❌ [{idx}] {nama}: {result[:200]}")

            progress.progress(100, text="Selesai!")

            with log_area:
                for line in log_lines:
                    icon = "🟢" if line.startswith("✅") else "🔴"
                    st.markdown(f"`{line}`")

            # ── Statistik singkat ─────────────────────────────
            st.divider()
            if results:
                st.success(f"✅ {len(results)} file berhasil di-generate!")
            if errors:
                st.error(f"❌ {len(errors)} file gagal: {', '.join(errors)}")

            # ── Metrics dari semua item ───────────────────────
            all_items = []
            for its in all_items_map.values(): all_items.extend(its)

            if all_items:
                STATUS_LAKU = ("SOLD","TERJUAL","LAKU")
                n_sold   = sum(1 for i in all_items if i["status"] in STATUS_LAKU)
                n_tap    = sum(1 for i in all_items if resolve_section(i["status"]) == "TAP")
                n_limit  = sum(1 for i in all_items if resolve_section(i["status"]) == "LIMIT")
                lim_val  = sum(i.get("limit",0) or 0 for i in all_items)
                sold_val = sum(i.get("harga_terbentuk",0) or 0
                               for i in all_items if i["status"] in STATUS_LAKU)

                pct_sold  = n_sold / len(all_items) * 100 if all_items else 0
                pct_harga = sold_val / lim_val * 100 if lim_val else 0
                c1, c2, c3, c4, c5 = st.columns(5)
                c1.metric("Total Lot",       len(all_items))
                c2.metric("Terjual (SOLD)",  n_sold,
                           delta=f"↑ {pct_sold:.0f}%")
                c3.metric("TAP",             n_tap)
                c4.metric("Total Limit",     _fmt_currency(lim_val))
                c5.metric("Harga Terbentuk", _fmt_currency(sold_val),
                           delta=f"↑ {pct_harga:.1f}%")

            # ── Download Buttons ──────────────────────────────
            st.divider()
            st.markdown("### 📥 Download Hasil")
            if len(results) == 1:
                out_name, doc_bytes = results[0]
                st.download_button(
                    label=f"⬇️ Download {out_name}",
                    data=doc_bytes, file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, type="primary")
            elif len(results) > 1:
                # Kemas semua jadi ZIP
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for out_name, doc_bytes in results:
                        zf.writestr(out_name, doc_bytes)
                zip_buf.seek(0)
                zip_name = f"RISALAH_BATCH_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
                st.download_button(
                    label=f"⬇️ Download Semua ({len(results)} file) sebagai ZIP",
                    data=zip_buf.getvalue(), file_name=zip_name,
                    mime="application/zip",
                    use_container_width=True, type="primary")

                st.markdown("**atau download satu per satu:**")
                cols = st.columns(min(len(results), 3))
                for i, (out_name, doc_bytes) in enumerate(results):
                    with cols[i % 3]:
                        st.download_button(
                            label=f"⬇️ {out_name[:35]}...",
                            data=doc_bytes, file_name=out_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_{i}", use_container_width=True)

            # Simpan ke session state untuk tab Statistik
            if all_items:
                st.session_state["all_items"]     = all_items
                st.session_state["all_items_map"] = all_items_map
                st.session_state["meta"]          = meta

    # ╔══════════════════════════════════════════════════════════════
    # ║  TAB 2 — STATISTIK LELANG
    # ╚══════════════════════════════════════════════════════════════
    with tab_stat:
        all_items = st.session_state.get("all_items")
        meta_stat = st.session_state.get("meta", meta)

        if not all_items:
            # Coba baca langsung dari upload jika belum generate
            if uploaded_files:
                preview_items = []
                for upf in uploaded_files:
                    upf.seek(0)
                    try:
                        it = read_excel_data(upf, sheet_name)
                        preview_items.extend(it)
                    except: pass
                all_items = preview_items if preview_items else None

        if not all_items:
            st.info("Upload file Excel dan jalankan Generate terlebih dahulu untuk melihat statistik.", icon="📊")
            st.stop()

        st.markdown("### 📊 Statistik Lelang")

        # ── Pilih file jika banyak ────────────────────────────
        items_map = st.session_state.get("all_items_map")
        if items_map and len(items_map) > 1:
            options = ["Semua File"] + list(items_map.keys())
            sel = st.selectbox("Filter File:", options)
            if sel != "Semua File":
                all_items = items_map[sel]
                display_name = sel
            else:
                display_name = "Semua File"
        else:
            display_name = (uploaded_files[0].name if uploaded_files else "Data")

        # ── Hitung Statistik ─────────────────────────────────
        fig_pie, fig_bar, stat = build_stat_charts(all_items, display_name)

        # ── Metrics ── pakai st.metric native + _fmt_currency() agar tidak terpotong
        pct_sold_s  = stat["sold"]  / stat["total"] * 100 if stat["total"] else 0
        pct_harga_s = stat["pencapaian"]
        c1, c2, c3, c4, c5, c6 = st.columns(6)
        c1.metric("Total Lot",       stat["total"])
        c2.metric("Terjual (SOLD)",  stat["sold"],
                   delta=f"↑ {pct_sold_s:.0f}%")
        c3.metric("TAP",             stat["tap"])
        c4.metric("Not Sold",        stat["limit"])
        c5.metric("Total Limit",     _fmt_currency(stat["total_limit"]))
        c6.metric("Harga Terbentuk", _fmt_currency(stat["total_harga"]),
                   delta=f"↑ {pct_harga_s:.1f}%")

        st.divider()

        # ── Charts ────────────────────────────────────────────
        col_pie, col_bar = st.columns(2)
        with col_pie:
            st.pyplot(fig_pie, use_container_width=True)
        with col_bar:
            st.pyplot(fig_bar, use_container_width=True)

        plt.close("all")

        st.divider()

        # ── Tabel Detail ──────────────────────────────────────
        with st.expander("📋 Tabel Detail Semua Lot", expanded=False):
            STATUS_LAKU = ("SOLD","TERJUAL","LAKU")
            table_data = []
            for i in all_items:
                s = resolve_section(i["status"])
                table_data.append({
                    "LOT"     : i["lot"],
                    "NOPOL"   : i["nopol"],
                    "MERK"    : i["merk"],
                    "MODEL"   : i["model"],
                    "TAHUN"   : i["tahun"],
                    "STATUS"  : s,
                    "LIMIT"   : format_rupiah(i.get("limit",0)),
                    "TERBENTUK": format_rupiah(i.get("harga_terbentuk",0)) if i["status"] in STATUS_LAKU else "-",
                })
            st.dataframe(table_data, use_container_width=True, height=350)

        # ── Export PDF Statistik ──────────────────────────────
        st.markdown("### 📄 Export Laporan Performa (PDF)")
        st.caption("Satu halaman PDF berisi tabel statistik + grafik pie & bar.")

        if st.button("🔄 Generate PDF Statistik", use_container_width=True):
            with st.spinner("Membuat PDF..."):
                fig_pie2, fig_bar2, _ = build_stat_charts(all_items, display_name)
                pdf_bytes = build_stat_pdf(all_items, meta_stat, display_name, fig_pie2, fig_bar2)
                plt.close("all")

            pdf_name = f"STATISTIK_LELANG_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
            st.download_button(
                label=f"⬇️ Download {pdf_name}",
                data=pdf_bytes, file_name=pdf_name,
                mime="application/pdf",
                use_container_width=True, type="primary")
            st.success("✅ PDF statistik siap diunduh!")


if __name__ == "__main__":
    main()
